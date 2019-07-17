# coding=utf-8
# pzw
# 20190716
# v0.2 完成字符串替换，完成图片插入，完成表格中的字符串替换


from docx import Document

## 字符串替换功能
def replaceParagraphString(document, tag, replaceString):
	paragraphs = document.paragraphs
	for p in paragraphs:
		if tag in p.text:
			for r in p.runs:
				if tag in r.text:
					r.text = replaceString

## 表格中的字符串替换功能
def replaceTableString(document, tag, replaceString):
	tables = document.tables
	for t in tables:
		rows = t.rows
		for r in rows:
			cells = r.cells
			for c in cells:
				paragraphs = c.paragraphs
				for p in paragraphs:
					if tag in p.text:
						for run in p.runs:
							if tag in run.text:
								run.text = replaceString

## 图片插入功能
def insertPicture(document, tag, picturePath):
	for p in document.paragraphs:
		if tag in p.text:
			replaceParagraphString(document, tag, " ")
			run = p.runs[0]
			if "(" in tag and ")" in tag:
				width = int(tag.split("(")[1].split(",")[0])
				height = int(tag.split(")")[0].split(",")[1])
				run.add_picture(picturePath, width*100000, height*100000)
			else:
				run.add_picture(picturePath)


## 表格插入功能
def fillTable(document, tag, insertTable):
	pass



## 测试脚本
docxFile = Document("test.docx")
# replaceTableString(docxFile, "#[CHARACTER-2]#", "xxfffxx")
insertPicture(docxFile, "#[PICTURE-(20,20)]#", "3e5cb4629d991c4b826788fe3e2877e5_hd.jpg")
replaceParagraphString(docxFile, "#[PARAGRAPH-1]#", "xxxxxx")
replaceTableString(docxFile, "#[CHARACTER-2]#", "xxxxxsga")
docxFile.save("test3.docx")