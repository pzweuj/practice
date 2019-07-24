# coding=utf-8
# pzw
# 20190722
# v0.8 支持每个节定义不同的页眉页脚
# v0.7 删除表格空行
# v0.6 去除冗余代码
# v0.5 修复表格第一行不能是纯数字的bug
# v0.4 页眉页脚
# v0.3 表格的插入，超链接的插入
# v0.2 完成字符串替换，完成图片插入，完成表格中的字符串替换

from docx import Document
from docx.oxml.shared import OxmlElement
from docx.oxml.shared import qn
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.opc.constants import RELATIONSHIP_TYPE
import pandas as pd


## 超链接
# 功能是在一个段落后增加超链接，未找到文本替换的方法
# 参考 https://stackoverflow.com/questions/47666642/adding-an-hyperlink-in-msword-by-using-python-docx
def add_hyperlink(paragraph, text, url):
	part = paragraph.part
	r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
	hyperlink = OxmlElement("w:hyperlink")
	hyperlink.set(qn("r:id"), r_id, )
	new_run = OxmlElement("w:r")
	rPr = OxmlElement("w:rPr")
	new_run.append(rPr)
	new_run.text = text
	hyperlink.append(new_run)
	r = paragraph.add_run()
	r._r.append(hyperlink)
	r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
	r.font.underline = True
	return hyperlink

## 字符串替换功能
def replaceParagraphString(document, tag, replaceString):
	paragraphs = document.paragraphs
	for p in paragraphs:
		if tag in p.text:
			for r in p.runs:
				if tag in r.text:
					r.text = unicode(replaceString, "utf-8")


## 表格中的字符串替换功能
def replaceTableString(document, tag, replaceString):
	tables = document.tables
	for t in tables:
		rows = t.rows
		for r in rows:
			cells = r.cells
			for c in cells:
				paragraphs = c.paragraphs
				for p in paragraphs:
					if tag in p.text:
						for run in p.runs:
							if tag in run.text:
								run.text = unicode(replaceString, "utf-8")

## 图片插入功能
def insertPicture(document, tag, picturePath):
	for p in document.paragraphs:
		if tag in p.text:
			replaceParagraphString(document, tag, " ")
			run = p.runs[0]
			if "(" in tag and ")" in tag:
				width = int(tag.split("(")[1].split(",")[0])
				height = int(tag.split(")")[0].split(",")[1])
				run.add_picture(picturePath, width*100000, height*100000)
			else:
				run.add_picture(picturePath)


## 表格初始化
def OriginTableReadyToFill(tableFile):
	table = pd.read_csv(tableFile, header=None, sep="\t")
	return table

## 删除表格行
def remove_row(table, row):
	tbl = table._tbl
	tr = row._tr
	tbl.remove(tr)

## 表格插入功能
def fillTable(document, tag, insertTable):
	tableToFill = OriginTableReadyToFill(insertTable)
	rowToFill = len(tableToFill)
	columnToFill = len(tableToFill.columns)

	for t in range(len(document.tables)):
		rows = document.tables[t].rows
		for r in range(len(rows)):
			cells = rows[r].cells
			for c in range(len(cells)):
				text = cells[c].text
				if tag in text:
					table_id = t
					row_id = r
					cell_id = c

	table = document.tables[table_id]

	# 格式刷
	cellAlignmentList = []
	for cell in table.rows[row_id].cells:
		cellAlignmentList.append(cell.vertical_alignment)

	styleList = []
	for cell in table.rows[row_id].cells:
		styleList.append(cell.paragraphs[0].style)

	alignmentList = []
	for cell in table.rows[row_id].cells:
		alignmentList.append(cell.paragraphs[0].alignment)

	boldList = []
	for cell in table.rows[row_id].cells:
		boldList.append(cell.paragraphs[0].runs[0].bold)

	italicList = []
	for cell in table.rows[row_id].cells:
		italicList.append(cell.paragraphs[0].runs[0].italic)

	fontNameList = []
	for cell in table.rows[row_id].cells:
		fontNameList.append(cell.paragraphs[0].runs[0].font.name)

	fontSizeList = []
	for cell in table.rows[row_id].cells:
		fontSizeList.append(cell.paragraphs[0].runs[0].font.size)

	colorList = []
	for cell in table.rows[row_id].cells:
		colorList.append(cell.paragraphs[0].runs[0].font.color.rgb)

	highlight_colorList = []
	for cell in table.rows[row_id].cells:
		highlight_colorList.append(cell.paragraphs[0].runs[0].font.highlight_color)

	# 判断行数是否足够，如果不够就添加
	if len(table.rows) - row_id < rowToFill:
		addRowAmount = rowToFill - len(table.rows) + row_id
		for i in range(addRowAmount):
			table.add_row()

	# 填充内容
	start = 0
	while row_id <= rowToFill:
		for co in range(columnToFill):
			table.cell(row_id, co).text = unicode(str(tableToFill.iloc[start, co]), "utf-8")
			table.cell(row_id, co).paragraphs[0].style = styleList[co]
			table.cell(row_id, co).paragraphs[0].alignment = alignmentList[co]

			for r in table.cell(row_id, co).paragraphs[0].runs:
				r.bold = boldList[co]
				r.italic = italicList[co]
				r.font.name = fontNameList[co]
				r.font.size = fontSizeList[co]
				r.font.color.rgb = colorList[co]
				r.font.highlight_color = highlight_colorList[co]

			table.cell(row_id, co).vertical_alignment = cellAlignmentList[co]

		start += 1
		row_id += 1

	# 删除表格空行
	for row in table.rows:
		pString = ""
		for cell in row.cells:
			for p in cell.paragraphs:
				pString = pString + p.text

		if pString == "":
			remove_row(table, row)

	del tableToFill


## 页脚
def footer(document, tag, replaceString):
	for s in document.sections:
		footer = s.footer
		for p in footer.paragraphs:
			for r in p.runs:
				if tag in r.text:
					r.text = unicode(replaceString, "utf-8")

## 页眉
def header(document, tag, replaceString):
	for s in document.sections:
		header = s.header
		for p in header.paragraphs:
			for r in p.runs:
				if tag in r.text:
					r.text = unicode(replaceString, "utf-8")



## 函数合并
def WordWriter(inputDocx, outputDocx, replaceDict):
	document = Document(inputDocx)
	for i in replaceDict:
		if "#[TABLE" in i:
			print(i)
			fillTable(document, i, replaceDict[i])

		elif "#[IMAGE" in i:
			print(i)
			insertPicture(document, i, replaceDict[i])

		elif "#[TBS" in i:
			print(i)
			replaceTableString(document, i, replaceDict[i])

		elif "#[FOOTER" in i:
			print(i)
			footer(document, i, replaceDict[i])

		elif "#[HEADER" in i:
			print(i)
			header(document, i, replaceDict[i])

		else:
			replaceParagraphString(document, i, replaceDict[i])
	document.save(outputDocx)

## 测试脚本
# testDict = {}
# testDict["#[HEADER-1]#"] = "模板测试"
# testDict["#[HEADER-2]#"] = "2019年7月18日"
# testDict["#[NAME]#"] = "测试模板"
# testDict["#[fullParagraph]#"] = "这是一段测试段落，通过WordWriter输入。"
# testDict["#[TBS-1]#"] = "未突变"
# testDict["#[FOOTER]#"] = "页脚测试"

# # 此处输入的是文件路径
# testDict["#[TABLE-1]#"] = "testTable.txt"
# testDict["#[IMAGE-1-(30,30)]#"] = "testPicture.png"
# testDict["#[IMAGE-2]#"] = "testPicture.png"

# # 使用主函数进行报告填充
# WordWriter("test.docx", "testOut.docx", testDict)